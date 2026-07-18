import os
import time
from pathlib import Path
from dotenv import load_dotenv
from groq import Groq
from pydantic import BaseModel, Field
import json
from pypdf import PdfReader
from docx import Document

load_dotenv()
my_api_key = os.getenv("GROQ_API_KEY")

if not my_api_key:
    raise ValueError("API key kaha hai bhai")

client = Groq(api_key=my_api_key)
model = "openai/gpt-oss-120b"

# ---------------- JOB DESCRIPTION ----------------
job_description = """ ... (your Amazon JD text here) ... """

class JobD(BaseModel):
    role: str
    required_skills: list[str]
    preferred_skills: list[str]
    minimum_experience: float | None
    education_requirements: list[str]
    responsibilities: list[str]

jobd_schema = JobD.model_json_schema()

system_prompt = f"""
You are an expert HR assistant.
Return ONLY valid JSON matching this schema:
{jobd_schema}
"""

user_prompt = f"Analyze the following job description:\n{job_description}"

message_system = {"role": "system", "content": system_prompt}
message_user = {"role": "user", "content": user_prompt}
response_format = {"type": "json_object"}

messages = [message_system, message_user]
response = client.chat.completions.create(model=model, messages=messages, response_format=response_format)
answer = response.choices[0].message.content
job_data = json.loads(answer)
job = JobD(**job_data)

print("Job minimum experience:", job.minimum_experience)
print("Job education requirements:", job.education_requirements)

# ---------------- RESUME PARSER ----------------
class MatchResult(BaseModel):
    score: float
    details: dict

class Experience(BaseModel):
    company: str | None = None
    role: str | None = None
    duration: str | None = None
    description: str | None = None
    skills_used: list[str] = []

class Resume(BaseModel):
    name: str | None = None
    email: str | None = None
    phone: str | None = None
    total_experience_years: float | None = None
    skills: list[str] = []
    experiences: list[Experience] = []
    education: list[str] = []
    projects: list[str] = []
    certifications: list[str] = []

resume_schema = Resume.model_json_schema()

def final_score(job, resume):
    match_schema = MatchResult.model_json_schema()
    prompt = f"""
    You are an HR recruiter.
    Compare the candidate's resume with the job description.
    JOB DESCRIPTION:
    {job.model_dump_json(indent=2)}
    CANDIDATE RESUME:
    {resume.model_dump_json(indent=2)}
    Return JSON matching this schema:
    {match_schema}
    """
    message = {"role": "user", "content": prompt}
    response = client.chat.completions.create(model=model, messages=[message], response_format={"type":"json_object"})
    data = json.loads(response.choices[0].message.content)
    return MatchResult(**data)

def parse_resume(resume_text):
    system_prompt = f"""
    You are an expert resume parser.
    Return ONLY valid JSON matching this schema:
    {resume_schema}
    """
    user_prompt = f"Parse the following resume:\n{resume_text}"
    messages = [{"role":"system","content":system_prompt},{"role":"user","content":user_prompt}]
    response = client.chat.completions.create(model=model, messages=messages, response_format={"type":"json_object"})
    raw_output = response.choices[0].message.content
    data = json.loads(raw_output)
    return Resume(**data)

# ---------------- FILE READERS ----------------
def read_pdf(file_path):
    reader = PdfReader(file_path)
    text = ""
    for page in reader.pages:
        page_text = page.extract_text()
        if page_text:
            text += page_text + "\n"
    return text

def read_docx(file_path):
    document = Document(file_path)
    text = ""
    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            text += paragraph.text + "\n"
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text += cell.text + "\n"
    return text

def read_resume(file_path):
    if file_path.suffix.lower() == ".pdf":
        return read_pdf(file_path)
    elif file_path.suffix.lower() == ".docx":
        return read_docx(file_path)
    else:
        return None

# ---------------- MAIN LOOP ----------------
resume_folder = Path("C:/Users/yuvan/padho-ai-engineer/week1/day5/resumes")

all_results = []
for file_path in resume_folder.iterdir():
    if file_path.suffix.lower() not in [".pdf", ".docx"]:
        continue
    print("\nProcessing:", file_path.name)
    resume_text = read_resume(file_path)
    if not resume_text:
        print("⚠️ Could not read:", file_path.name)
        continue
    parsed_resume = parse_resume(resume_text)
    time.sleep(5)
    result = final_score(job, parsed_resume)
    time.sleep(5)
    print("Score:", result.score)
    all_results.append({"name": parsed_resume.name, "score": result.score, "details": result.details})

all_results.sort(key=lambda c: c["score"], reverse=True)
top_2 = all_results[:2]
worst_2 = all_results[-2:]

print("\nTOP 2 CANDIDATES")
for c in top_2:
    print(c["name"], "-", c["score"], "%")
    print(c["details"])

print("\nLOWEST 2 CANDIDATES")
for c in worst_2:
    print(c["name"], "-", c["score"], "%")
    print(c["details"])
